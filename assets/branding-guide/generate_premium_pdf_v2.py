#!/usr/bin/env python3
"""
Generate Premium PDF Brand Guide for The Right Path Educational Consulting Inc.
Version 2 - Fixed style conflicts and image generation

This script:
1. Extracts content from the Word document
2. Uses existing hero images + creates branded graphics with PIL
3. Builds a premium PDF with ReportLab
"""

import os
import sys
from pathlib import Path
from datetime import datetime

# PDF Generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, white, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.platypus.flowables import HRFlowable

# Word Document Processing
from docx import Document

# Image Processing
from PIL import Image as PILImage, ImageDraw, ImageFont

# ==================== CONFIGURATION ====================

# Paths
BASE_DIR = Path(r"C:\Users\MarieLexisDad\The Right Path\assets\branding-guide")
WORD_DOC = BASE_DIR / "The_Right_Path_Brand_Guide_COMPLETE.docx"
OUTPUT_PDF = BASE_DIR / "The_Right_Path_Brand_Guide_PREMIUM_v2.pdf"
ASSETS_DIR = BASE_DIR / "pdf_assets"

# Create assets directory
ASSETS_DIR.mkdir(exist_ok=True)

# Brand Colors (RGB tuples for PIL, HexColor for ReportLab)
NAVY_RGB = (11, 29, 58)
GOLD_RGB = (255, 211, 58)
ELECTRIC_BLUE_RGB = (0, 217, 255)
WHITE_RGB = (255, 255, 255)
DARK_GRAY_RGB = (51, 51, 51)

NAVY = HexColor('#0B1D3A')
GOLD = HexColor('#FFD33A')
ELECTRIC_BLUE = HexColor('#00D9FF')
DARK_GRAY = HexColor('#333333')
LIGHT_GRAY = HexColor('#F5F5F5')

# ==================== BRANDED GRAPHICS GENERATION WITH PIL ====================

def create_cover_graphic(output_path, width=2550, height=3300):
    """Create a branded cover page graphic with elegant diagonal design."""
    print(f"  Creating cover graphic: {Path(output_path).name}")
    import math

    img = PILImage.new('RGB', (width, height), NAVY_RGB)
    draw = ImageDraw.Draw(img)

    # Create elegant diagonal stripe pattern (top-left to bottom-right)
    stripe_width = 200
    stripe_gap = 400

    # Draw multiple diagonal gold accent stripes
    for offset in range(-height, width + height, stripe_gap):
        points = [
            (offset, 0),
            (offset + stripe_width, 0),
            (offset + stripe_width + height, height),
            (offset + height, height)
        ]
        # Subtle navy variation stripe
        draw.polygon(points, fill=(20, 40, 75))

    # Large gold diagonal band (dominant visual element)
    gold_band_offset = width // 3
    gold_band_points = [
        (gold_band_offset, 0),
        (gold_band_offset + 400, 0),
        (gold_band_offset + 400 + height * 0.6, height),
        (gold_band_offset + height * 0.6, height)
    ]
    draw.polygon(gold_band_points, fill=GOLD_RGB)

    # Electric blue accent stripe (parallel to gold)
    blue_offset = gold_band_offset + 420
    blue_points = [
        (blue_offset, 0),
        (blue_offset + 30, 0),
        (blue_offset + 30 + height * 0.6, height),
        (blue_offset + height * 0.6, height)
    ]
    draw.polygon(blue_points, fill=ELECTRIC_BLUE_RGB)

    # Create elegant curved element at bottom
    for i in range(80):
        y = height - 200 + i
        curve_height = int(40 * math.sin(i * math.pi / 80))
        draw.line([(0, y - curve_height), (width, y - curve_height)],
                  fill=GOLD_RGB if i < 40 else ELECTRIC_BLUE_RGB, width=2)

    # Add horizontal gold bar at very bottom
    draw.rectangle([(0, height - 60), (width, height)], fill=GOLD_RGB)

    # Add corner accent triangles
    # Top-left corner
    draw.polygon([(0, 0), (200, 0), (0, 200)], fill=GOLD_RGB)

    # Bottom-right corner
    draw.polygon([(width, height), (width - 200, height), (width, height - 200)], fill=GOLD_RGB)

    # Central focal circle with electric blue ring
    center_x, center_y = width // 2 + 300, height // 2
    radius = 350
    # Outer gold ring
    for r in range(radius + 20, radius + 30):
        draw.ellipse([(center_x - r, center_y - r), (center_x + r, center_y + r)],
                     outline=GOLD_RGB, width=3)
    # Electric blue ring
    for r in range(radius + 5, radius + 15):
        draw.ellipse([(center_x - r, center_y - r), (center_x + r, center_y + r)],
                     outline=ELECTRIC_BLUE_RGB, width=2)
    # Inner navy circle
    draw.ellipse([(center_x - radius, center_y - radius),
                  (center_x + radius, center_y + radius)],
                 fill=(15, 35, 70), outline=GOLD_RGB, width=4)

    img.save(output_path, quality=95)
    print(f"    [OK] Created: {output_path}")
    return output_path

def create_logo_concept(output_path, variant=1, size=1024):
    """Create a branded logo concept graphic."""
    print(f"  Creating logo concept {variant}: {Path(output_path).name}")

    img = PILImage.new('RGB', (size, size), WHITE_RGB)
    draw = ImageDraw.Draw(img)

    center = size // 2
    hex_radius = size // 3

    # Draw hexagon
    import math
    points = []
    for i in range(6):
        angle = math.pi / 3 * i - math.pi / 2
        px = center + hex_radius * math.cos(angle)
        py = center + hex_radius * math.sin(angle)
        points.append((px, py))

    if variant == 1:
        # Primary - filled navy with gold accent
        draw.polygon(points, fill=NAVY_RGB, outline=GOLD_RGB)
        # Inner hexagon
        inner_points = []
        for i in range(6):
            angle = math.pi / 3 * i - math.pi / 2
            px = center + hex_radius * 0.6 * math.cos(angle)
            py = center + hex_radius * 0.6 * math.sin(angle)
            inner_points.append((px, py))
        draw.polygon(inner_points, fill=GOLD_RGB)

    elif variant == 2:
        # Secondary - outline style
        draw.polygon(points, fill=None, outline=NAVY_RGB)
        for i, p in enumerate(points):
            next_p = points[(i + 1) % 6]
            draw.line([p, next_p], fill=NAVY_RGB, width=8)
        # Gold accent
        draw.polygon(points, fill=None, outline=GOLD_RGB)

    else:
        # Icon - gradient effect
        draw.polygon(points, fill=NAVY_RGB)
        # Electric blue inner
        inner_points = []
        for i in range(6):
            angle = math.pi / 3 * i - math.pi / 2
            px = center + hex_radius * 0.5 * math.cos(angle)
            py = center + hex_radius * 0.5 * math.sin(angle)
            inner_points.append((px, py))
        draw.polygon(inner_points, fill=ELECTRIC_BLUE_RGB)

    img.save(output_path, quality=95)
    print(f"    [OK] Created: {output_path}")
    return output_path

def create_color_palette_visual(output_path, width=2400, height=600):
    """Create color palette visualization."""
    print(f"  Creating color palette: {Path(output_path).name}")

    img = PILImage.new('RGB', (width, height), WHITE_RGB)
    draw = ImageDraw.Draw(img)

    colors = [
        (NAVY_RGB, "Primary: Navy", "#0B1D3A"),
        (GOLD_RGB, "Secondary: Gold", "#FFD33A"),
        (ELECTRIC_BLUE_RGB, "Accent: Electric Blue", "#00D9FF")
    ]

    swatch_width = width // 3
    for i, (color, name, hex_code) in enumerate(colors):
        x = i * swatch_width
        # Color swatch
        draw.rectangle([(x + 20, 20), (x + swatch_width - 20, height - 80)],
                       fill=color)
        # Hex code text area
        text_color = WHITE_RGB if color == NAVY_RGB else DARK_GRAY_RGB
        # Simple text overlay (would use font in production)
        draw.rectangle([(x + 20, height - 70), (x + swatch_width - 20, height - 20)],
                       fill=(240, 240, 240))

    img.save(output_path, quality=95)
    print(f"    [OK] Created: {output_path}")
    return output_path

def create_social_template(output_path, template_type="instagram", size=1080):
    """Create social media template graphic."""
    print(f"  Creating {template_type} template: {Path(output_path).name}")

    if template_type in ["facebook", "linkedin", "email"]:
        # Horizontal format
        img = PILImage.new('RGB', (1200, 630), NAVY_RGB)
    else:
        # Square format
        img = PILImage.new('RGB', (size, size), NAVY_RGB)

    draw = ImageDraw.Draw(img)
    w, h = img.size

    if template_type == "podcast":
        # Gold frame border
        draw.rectangle([(30, 30), (w-30, h-30)], outline=GOLD_RGB, width=8)
        # Central content area
        draw.rectangle([(60, 60), (w-60, h-60)], fill=(20, 40, 70))
        # Episode title area
        draw.rectangle([(100, h//2 - 100), (w-100, h//2 + 100)], fill=GOLD_RGB)

    elif template_type == "instagram_quote":
        # Quote card style
        draw.rectangle([(50, 200), (w-50, h-200)], fill=GOLD_RGB)
        draw.rectangle([(70, 220), (w-70, h-220)], fill=NAVY_RGB)
        # Quote marks placeholder
        draw.ellipse([(100, 260), (180, 340)], fill=ELECTRIC_BLUE_RGB)

    elif template_type == "instagram_stat":
        # Stats infographic
        draw.rectangle([(0, 0), (w, h//3)], fill=GOLD_RGB)
        draw.rectangle([(0, h//3), (w, h//3 + 10)], fill=ELECTRIC_BLUE_RGB)
        # Placeholder for big number
        draw.ellipse([(w//2 - 150, h//2), (w//2 + 150, h//2 + 300)], fill=(30, 50, 80))

    elif template_type == "facebook":
        # Horizontal banner
        draw.rectangle([(0, 0), (w, 100)], fill=GOLD_RGB)
        draw.rectangle([(0, 100), (w, 110)], fill=ELECTRIC_BLUE_RGB)
        # Content area
        draw.rectangle([(50, 150), (w//2, h-50)], fill=(20, 40, 70))

    elif template_type == "linkedin":
        # Professional layout
        draw.rectangle([(0, 0), (80, h)], fill=GOLD_RGB)
        draw.rectangle([(80, 0), (90, h)], fill=ELECTRIC_BLUE_RGB)
        # Content area
        draw.rectangle([(120, 50), (w-50, h-50)], fill=(20, 40, 70))

    else:  # email header
        # Horizontal banner
        draw.rectangle([(0, 0), (w, h//3)], fill=GOLD_RGB)
        draw.rectangle([(0, h//3), (w, h//3 + 8)], fill=ELECTRIC_BLUE_RGB)

    img.save(output_path, quality=95)
    print(f"    [OK] Created: {output_path}")
    return output_path

# ==================== CONTENT EXTRACTION ====================

def extract_word_content():
    """Extract content from Word document."""
    print("\n[Phase 1] Extracting content from Word document...")

    doc = Document(str(WORD_DOC))

    content = {
        'paragraphs': [],
        'tables': []
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Clean "PAGE X -" prefixes
        if text.startswith("PAGE "):
            parts = text.split(" - ", 1)
            if len(parts) > 1:
                text = parts[1]

        style_name = para.style.name if para.style else "Normal"

        content['paragraphs'].append({
            'text': text,
            'style': style_name,
            'is_heading': 'Heading' in style_name or (text.isupper() and len(text) > 3 and len(text) < 50)
        })

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            content['tables'].append(table_data)

    print(f"  Extracted {len(content['paragraphs'])} paragraphs, {len(content['tables'])} tables")
    return content

# ==================== PDF GENERATION ====================

class BrandGuidePDF:
    """Premium PDF generator for The Right Path Brand Guide."""

    def __init__(self):
        self.styles = self._create_styles()
        self.story = []

    def _create_styles(self):
        """Create custom paragraph styles."""
        styles = getSampleStyleSheet()

        # Use TRP prefix to avoid conflicts with default styles
        custom_styles = {
            'TRPCoverTitle': ParagraphStyle(
                name='TRPCoverTitle',
                fontName='Helvetica-Bold',
                fontSize=36,
                textColor=white,
                alignment=TA_CENTER,
                spaceAfter=20
            ),
            'TRPCoverSubtitle': ParagraphStyle(
                name='TRPCoverSubtitle',
                fontName='Helvetica',
                fontSize=18,
                textColor=white,
                alignment=TA_CENTER,
                spaceAfter=40
            ),
            'TRPSectionHeader': ParagraphStyle(
                name='TRPSectionHeader',
                fontName='Helvetica-Bold',
                fontSize=28,
                textColor=NAVY,
                alignment=TA_LEFT,
                spaceBefore=30,
                spaceAfter=20
            ),
            'TRPSubsectionHeader': ParagraphStyle(
                name='TRPSubsectionHeader',
                fontName='Helvetica-Bold',
                fontSize=16,
                textColor=NAVY,
                alignment=TA_LEFT,
                spaceBefore=15,
                spaceAfter=10
            ),
            'TRPBodyText': ParagraphStyle(
                name='TRPBodyText',
                fontName='Helvetica',
                fontSize=11,
                textColor=DARK_GRAY,
                alignment=TA_JUSTIFY,
                leading=16,
                spaceBefore=4,
                spaceAfter=8
            ),
            'TRPQuoteText': ParagraphStyle(
                name='TRPQuoteText',
                fontName='Helvetica-Oblique',
                fontSize=13,
                textColor=NAVY,
                alignment=TA_CENTER,
                leading=18,
                leftIndent=30,
                rightIndent=30,
                spaceBefore=15,
                spaceAfter=15
            ),
            'TRPTableHeader': ParagraphStyle(
                name='TRPTableHeader',
                fontName='Helvetica-Bold',
                fontSize=10,
                textColor=white,
                alignment=TA_LEFT
            ),
            'TRPTableCell': ParagraphStyle(
                name='TRPTableCell',
                fontName='Helvetica',
                fontSize=9,
                textColor=DARK_GRAY,
                alignment=TA_LEFT,
                leading=12
            ),
            'TRPTOCItem': ParagraphStyle(
                name='TRPTOCItem',
                fontName='Helvetica',
                fontSize=12,
                textColor=NAVY,
                alignment=TA_LEFT,
                spaceBefore=8,
                spaceAfter=4,
                leftIndent=20
            )
        }

        for name, style in custom_styles.items():
            styles.add(style)

        return styles

    def add_cover_page(self, cover_image_path=None):
        """Add cover page with image above text."""
        if cover_image_path and Path(cover_image_path).exists():
            # Large image above text (fit within margins, leave room for title)
            img = Image(cover_image_path, width=7*inch, height=4.5*inch)
            self.story.append(img)
            self.story.append(Spacer(1, 0.3*inch))
        else:
            # Text-based cover
            self.story.append(Spacer(1, 2*inch))

        # Text content
        self.story.append(Paragraph(
            "THE RIGHT PATH",
            self.styles['TRPCoverTitle']
        ))
        self.story.append(Paragraph(
            "EDUCATIONAL CONSULTING INC.",
            self.styles['TRPCoverSubtitle']
        ))
        self.story.append(Spacer(1, 0.3*inch))
        self.story.append(Paragraph(
            "Brand Guidelines",
            ParagraphStyle(
                'TRPBrandGuide',
                fontName='Helvetica-Bold',
                fontSize=24,
                textColor=GOLD,
                alignment=TA_CENTER
            )
        ))
        self.story.append(Spacer(1, 0.5*inch))
        self.story.append(Paragraph(
            "Demystifying AI for Black and Latino Educators",
            ParagraphStyle(
                'TRPTagline',
                fontName='Helvetica-Oblique',
                fontSize=14,
                textColor=ELECTRIC_BLUE,
                alignment=TA_CENTER
            )
        ))

        self.story.append(PageBreak())

    def add_toc(self):
        """Add table of contents."""
        self.story.append(Paragraph("TABLE OF CONTENTS", self.styles['TRPSectionHeader']))
        self.story.append(HRFlowable(width="30%", thickness=3, color=GOLD, spaceAfter=20))

        toc_items = [
            "Mission & Vision",
            "Brand Story",
            "Target Audience",
            "Brand Voice & Tone",
            "Logo System",
            "Color Palette",
            "Typography",
            "Imagery Guidelines",
            "Brand Pillars",
            "Visual Language",
            "Application Examples",
            "Do's & Don'ts",
            "Social Media Guidelines",
            "Contact & Resources"
        ]

        for i, item in enumerate(toc_items, 1):
            self.story.append(Paragraph(f"{i}. {item}", self.styles['TRPTOCItem']))

        self.story.append(PageBreak())

    def add_section_header(self, title):
        """Add a section header with gold accent."""
        # Clean any PAGE X - prefix
        if title.upper().startswith("PAGE "):
            parts = title.split(" - ", 1)
            if len(parts) > 1:
                title = parts[1]

        self.story.append(Paragraph(title.upper(), self.styles['TRPSectionHeader']))
        self.story.append(HRFlowable(width="30%", thickness=3, color=GOLD, spaceBefore=5, spaceAfter=20))

    def add_subsection_header(self, title):
        """Add a subsection header."""
        self.story.append(Paragraph(title, self.styles['TRPSubsectionHeader']))

    def add_body_text(self, text):
        """Add body paragraph."""
        self.story.append(Paragraph(text, self.styles['TRPBodyText']))

    def add_quote(self, text):
        """Add styled quote."""
        self.story.append(Paragraph(f'"{text}"', self.styles['TRPQuoteText']))

    def add_image(self, image_path, width=6*inch, height=None, caption=None):
        """Add image with optional caption."""
        if Path(image_path).exists():
            if height:
                img = Image(image_path, width=width, height=height)
            else:
                img = Image(image_path, width=width)
            self.story.append(img)

            if caption:
                self.story.append(Paragraph(
                    caption,
                    ParagraphStyle(
                        'TRPImageCaption',
                        fontName='Helvetica-Oblique',
                        fontSize=9,
                        textColor=DARK_GRAY,
                        alignment=TA_CENTER,
                        spaceBefore=5,
                        spaceAfter=15
                    )
                ))
            else:
                self.story.append(Spacer(1, 15))

    def add_color_palette_section(self):
        """Add visual color palette section."""
        self.add_section_header("COLOR PALETTE")

        colors = [
            ("Primary: Navy Blue", "#0B1D3A", NAVY, "Authority, Trust, Professionalism"),
            ("Secondary: Gold", "#FFD33A", GOLD, "Optimism, Warmth, Excellence"),
            ("Accent: Electric Blue", "#00D9FF", ELECTRIC_BLUE, "Innovation, Technology, Future")
        ]

        for name, hex_code, color, meaning in colors:
            data = [
                [Paragraph(f"<b>{name}</b><br/>{hex_code}<br/><i>{meaning}</i>", self.styles['TRPTableCell'])]
            ]

            t = Table(data, colWidths=[6*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, 0), color),
                ('TEXTCOLOR', (0, 0), (0, 0), white if color == NAVY else black),
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (0, 0), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (0, 0), 20),
                ('BOTTOMPADDING', (0, 0), (0, 0), 20),
                ('BOX', (0, 0), (0, 0), 2, GOLD)
            ]))
            self.story.append(t)
            self.story.append(Spacer(1, 15))

    def add_typography_section(self):
        """Add typography section."""
        self.add_section_header("TYPOGRAPHY")

        self.add_subsection_header("Primary Font: Montserrat")
        self.story.append(Paragraph(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZ<br/>abcdefghijklmnopqrstuvwxyz<br/>0123456789",
            ParagraphStyle(
                'TRPFontSample1',
                fontName='Helvetica-Bold',
                fontSize=14,
                textColor=NAVY,
                leading=20,
                spaceBefore=10,
                spaceAfter=10
            )
        ))
        self.add_body_text("Used for: Headlines, Titles, Navigation, Call-to-Action buttons")

        self.story.append(Spacer(1, 20))

        self.add_subsection_header("Secondary Font: Open Sans")
        self.story.append(Paragraph(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZ<br/>abcdefghijklmnopqrstuvwxyz<br/>0123456789",
            ParagraphStyle(
                'TRPFontSample2',
                fontName='Helvetica',
                fontSize=14,
                textColor=DARK_GRAY,
                leading=20,
                spaceBefore=10,
                spaceAfter=10
            )
        ))
        self.add_body_text("Used for: Body text, Paragraphs, Descriptions, UI elements")

    def add_logo_section(self, logo_paths):
        """Add logo showcase section."""
        self.add_section_header("LOGO SYSTEM")
        self.add_body_text("Our logo system includes variations for different use cases while maintaining brand consistency across all touchpoints.")

        labels = ["Primary Logo", "Secondary Logo (Horizontal)", "Icon Logo"]
        for i, logo_path in enumerate(logo_paths[:3]):
            if Path(logo_path).exists():
                self.story.append(Spacer(1, 10))
                self.add_subsection_header(labels[i] if i < len(labels) else f"Logo Variation {i+1}")
                self.add_image(logo_path, width=2*inch, height=2*inch)
                # Add page break after each logo except the last
                if i < 2:
                    self.story.append(PageBreak())

    def add_templates_section(self, template_paths):
        """Add social media templates showcase."""
        self.add_section_header("APPLICATION EXAMPLES")
        self.add_body_text("These templates demonstrate how to apply the brand identity across various digital platforms.")

        for i, template_path in enumerate(template_paths):
            if Path(template_path).exists():
                name = Path(template_path).stem.replace("_", " ").title()
                self.add_subsection_header(name)
                self.add_image(template_path, width=3*inch, height=3*inch)
                # Page break after every 2 templates
                if (i + 1) % 2 == 0 and i < len(template_paths) - 1:
                    self.story.append(PageBreak())

    def add_table(self, table_data, col_widths=None):
        """Add a styled table."""
        if not table_data or len(table_data) < 2:
            return

        formatted_data = []
        for i, row in enumerate(table_data):
            formatted_row = []
            for cell in row:
                if i == 0:
                    formatted_row.append(Paragraph(str(cell)[:100], self.styles['TRPTableHeader']))
                else:
                    formatted_row.append(Paragraph(str(cell)[:200], self.styles['TRPTableCell']))
            formatted_data.append(formatted_row)

        if col_widths is None:
            num_cols = len(formatted_data[0])
            col_widths = [6.5*inch / num_cols] * num_cols

        t = Table(formatted_data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), NAVY),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BACKGROUND', (0, 1), (-1, -1), white),
            ('TEXTCOLOR', (0, 1), (-1, -1), DARK_GRAY),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, LIGHT_GRAY]),
            ('GRID', (0, 0), (-1, -1), 0.5, LIGHT_GRAY),
            ('BOX', (0, 0), (-1, -1), 1, NAVY),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))

        self.story.append(t)
        self.story.append(Spacer(1, 15))

    def add_contact_section(self):
        """Add contact and resources section."""
        self.add_section_header("CONTACT & RESOURCES")

        self.add_subsection_header("The Right Path Educational Consulting Inc.")
        contact_info = [
            "Website: therightpathedu.com",
            "Email: info@therightpathedu.com",
            "LinkedIn: linkedin.com/company/the-right-path-edu"
        ]
        for info in contact_info:
            self.add_body_text(info)

        self.story.append(Spacer(1, 20))

        self.add_subsection_header("Brand Assets")
        self.add_body_text("For brand assets, templates, and guidelines, please contact the marketing team. All brand materials should be used in accordance with these guidelines to maintain consistency and brand integrity.")

    def build(self, output_path):
        """Build the final PDF."""
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch
        )

        doc.build(self.story)
        print(f"\n  [OK] PDF saved: {output_path}")

# ==================== MAIN WORKFLOW ====================

def generate_graphics():
    """Generate all branded graphics using PIL."""
    print("\n[Phase 2] Generating branded graphics with PIL...")

    images = {
        'cover': None,
        'logos': [],
        'color_palette': None,
        'templates': []
    }

    # Cover graphic
    cover_path = str(ASSETS_DIR / "cover_graphic.png")
    create_cover_graphic(cover_path)
    images['cover'] = cover_path

    # Logo concepts
    for i in range(1, 4):
        logo_path = str(ASSETS_DIR / f"logo_concept_{i}.png")
        create_logo_concept(logo_path, variant=i)
        images['logos'].append(logo_path)

    # Color palette
    palette_path = str(ASSETS_DIR / "color_palette.png")
    create_color_palette_visual(palette_path)
    images['color_palette'] = palette_path

    # Social templates
    template_types = ["podcast", "instagram_quote", "instagram_stat", "facebook", "linkedin", "email"]
    for template_type in template_types:
        template_path = str(ASSETS_DIR / f"{template_type}_template.png")
        create_social_template(template_path, template_type)
        images['templates'].append(template_path)

    return images

def build_pdf(content, images):
    """Build the premium PDF."""
    print("\n[Phase 3] Building premium PDF...")

    pdf = BrandGuidePDF()

    # Cover page
    pdf.add_cover_page(images.get('cover'))

    # Table of contents
    pdf.add_toc()

    # Hero images from existing files
    hero_images = {
        'mission': BASE_DIR / 'page2_hero_image.png',
        'brand story': BASE_DIR / 'page3_hero_image.png',
        'target': BASE_DIR / 'page4_hero_image.png',
        'voice': BASE_DIR / 'page5_hero_image.png'
    }

    # Process content from Word doc
    current_section = None
    sections_added = set()
    table_idx = 0

    section_keywords = {
        'mission': 'MISSION & VISION',
        'vision': 'MISSION & VISION',
        'brand story': 'BRAND STORY',
        'target': 'TARGET AUDIENCE',
        'audience': 'TARGET AUDIENCE',
        'voice': 'BRAND VOICE & TONE',
        'tone': 'BRAND VOICE & TONE',
        'pillar': 'BRAND PILLARS',
        'imagery': 'IMAGERY GUIDELINES',
        'visual language': 'VISUAL LANGUAGE',
        'application': 'APPLICATION EXAMPLES',
        'do\'s': 'DO\'S & DON\'TS',
        'don\'t': 'DO\'S & DON\'TS',
        'social media': 'SOCIAL MEDIA GUIDELINES',
        'contact': 'CONTACT & RESOURCES'
    }

    for para in content['paragraphs']:
        text = para['text']
        is_heading = para['is_heading']
        text_lower = text.lower()

        # Check for new section
        for keyword, section_name in section_keywords.items():
            if keyword in text_lower and is_heading:
                if section_name not in sections_added:
                    if current_section:
                        pdf.story.append(PageBreak())
                    current_section = section_name
                    sections_added.add(section_name)
                    pdf.add_section_header(section_name)

                    # Add hero image if available (constrained size to fit page)
                    for hero_key, hero_path in hero_images.items():
                        if hero_key in text_lower and hero_path.exists():
                            pdf.add_image(str(hero_path), width=5*inch, height=3*inch)
                            break
                break
        else:
            # Add content
            if current_section:
                if is_heading and len(text) < 80:
                    pdf.add_subsection_header(text)
                elif text.startswith('"') or ('"' in text and len(text) < 200):
                    pdf.add_quote(text.strip('"\''))
                else:
                    pdf.add_body_text(text)

    # Add tables from content
    for table_data in content['tables'][:10]:  # Limit to first 10 tables
        if len(table_data) > 1:
            pdf.add_table(table_data)

    # Add custom sections
    pdf.story.append(PageBreak())
    pdf.add_logo_section(images['logos'])

    pdf.story.append(PageBreak())
    pdf.add_color_palette_section()

    pdf.story.append(PageBreak())
    pdf.add_typography_section()

    pdf.story.append(PageBreak())
    pdf.add_templates_section(images['templates'])

    pdf.story.append(PageBreak())
    pdf.add_contact_section()

    # Build PDF
    pdf.build(OUTPUT_PDF)
    return OUTPUT_PDF

def validate_pdf(pdf_path):
    """Validate the generated PDF."""
    print("\n[Phase 4] Validating PDF...")

    from PyPDF2 import PdfReader

    pdf = PdfReader(str(pdf_path))
    page_count = len(pdf.pages)
    file_size_mb = Path(pdf_path).stat().st_size / (1024 * 1024)

    print(f"  Pages: {page_count}")
    print(f"  File size: {file_size_mb:.2f} MB")

    issues = []
    if page_count < 10:
        issues.append(f"Page count ({page_count}) is less than expected")
    if file_size_mb > 20:
        issues.append(f"File size exceeds 20MB")

    if issues:
        print("  [!] Issues found:")
        for issue in issues:
            print(f"      - {issue}")
    else:
        print("  [OK] Validation passed!")

    return {'pages': page_count, 'size_mb': file_size_mb, 'valid': len(issues) == 0}

def main():
    """Main execution."""
    print("=" * 60)
    print("THE RIGHT PATH - PREMIUM BRAND GUIDE PDF GENERATOR v2")
    print("=" * 60)

    start = datetime.now()

    # Phase 1: Extract content
    content = extract_word_content()

    # Phase 2: Generate graphics
    images = generate_graphics()

    # Phase 3: Build PDF
    pdf_path = build_pdf(content, images)

    # Phase 4: Validate
    result = validate_pdf(pdf_path)

    elapsed = (datetime.now() - start).total_seconds()

    print("\n" + "=" * 60)
    print("COMPLETE")
    print("=" * 60)
    print(f"  Output: {pdf_path}")
    print(f"  Pages: {result['pages']}")
    print(f"  Size: {result['size_mb']:.2f} MB")
    print(f"  Time: {elapsed:.1f} seconds")
    print("=" * 60)

    return pdf_path

if __name__ == "__main__":
    main()
