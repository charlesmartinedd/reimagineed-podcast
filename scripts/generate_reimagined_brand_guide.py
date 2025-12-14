"""
ReimagineED Brand Guide Generator
Transforms AI in Action brand guide into executive-ready ReimagineED document.

Usage:
    python generate_reimagined_brand_guide.py
"""

import os
import sys
import time
import requests
from pathlib import Path
from io import BytesIO

# PDF processing
import fitz  # PyMuPDF

# Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from PIL import Image

# =============================================================================
# CONFIGURATION
# =============================================================================

# Paths
BASE_DIR = Path(r"C:\Users\MarieLexisDad\The Right Path")
PDF_SOURCE = BASE_DIR / "assets" / "AI-in-Action-Brand-Guide-Interactive.pdf"
STRATEGY_MD = BASE_DIR / "docs" / "reimagined-brand-strategy.md"
OUTPUT_DIR = BASE_DIR / "assets" / "branding-guide" / "reimagined"
EXTRACTED_IMAGES_DIR = BASE_DIR / "assets" / "branding-guide" / "extracted_images"
LOGOS_DIR = BASE_DIR / "assets" / "branding-guide" / "generated_logos"

# Brand Colors
COLORS = {
    'navy': RGBColor(11, 29, 58),        # #0B1D3A
    'gold': RGBColor(255, 211, 58),      # #FFD33A
    'electric_blue': RGBColor(0, 217, 255),  # #00D9FF
    'purple': RGBColor(123, 47, 255),    # #7B2FFF
    'coral': RGBColor(255, 107, 107),    # #FF6B6B
    'white': RGBColor(255, 255, 255),
    'charcoal': RGBColor(51, 51, 51),    # #333333
}

# =============================================================================
# IMAGE EXTRACTION
# =============================================================================

def extract_images_from_pdf(pdf_path: Path, output_dir: Path) -> dict:
    """Extract images from PDF, skipping page 6 (Leadership Team)."""
    output_dir.mkdir(parents=True, exist_ok=True)
    images = {}

    print(f"Extracting images from: {pdf_path}")
    doc = fitz.open(str(pdf_path))

    for page_num in range(len(doc)):
        # Skip Leadership Team page (page 6, index 5)
        if page_num == 5:
            print(f"  Skipping page {page_num + 1} (Leadership Team)")
            continue

        page = doc[page_num]
        image_list = page.get_images()

        if not image_list:
            continue

        images[page_num] = []

        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # Save image
                image_filename = f"page{page_num + 1}_img{img_index}.{image_ext}"
                image_path = output_dir / image_filename

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                images[page_num].append(str(image_path))
                print(f"  Extracted: {image_filename}")
            except Exception as e:
                print(f"  Error extracting image {img_index} from page {page_num + 1}: {e}")

    doc.close()
    return images


# =============================================================================
# LOGO GENERATION (Freepik API)
# =============================================================================

def generate_logo_freepik(output_path: Path) -> str:
    """Generate tech-forward ReimagineED logo via Freepik Mystic API."""
    api_key = os.getenv("FREEPIK_API_KEY")

    if not api_key:
        print("Warning: FREEPIK_API_KEY not set. Using placeholder for logo.")
        return None

    output_path.parent.mkdir(parents=True, exist_ok=True)

    prompt = """
    Modern tech-forward wordmark logo design for "ReimagineED" education technology brand.
    Clean minimalist typography with the word "REIMAGINE" in bold navy blue (#0B1D3A) geometric sans-serif font,
    followed by "ED" in bright gold (#FFD33A) with a subtle digital glow effect.
    Professional brand identity style, white background, suitable for corporate presentations.
    No icons or symbols, pure typography wordmark. High resolution, clean vector-style edges.
    Modern tech company aesthetic like Slack or Notion branding.
    """

    print("Generating logo via Freepik Mystic API...")

    try:
        # Create generation task
        response = requests.post(
            "https://api.freepik.com/v1/ai/mystic",
            headers={
                "x-freepik-api-key": api_key,
                "Content-Type": "application/json"
            },
            json={
                "prompt": prompt.strip(),
                "resolution": "2k",
                "aspect_ratio": "widescreen_16_9",
                "styling": {
                    "style": "photo"
                }
            },
            timeout=30
        )

        if response.status_code != 200:
            print(f"API Error: {response.status_code} - {response.text}")
            return None

        result = response.json()

        # Check if we got immediate results or need to poll
        if "data" in result and result["data"]:
            # Direct result
            image_url = result["data"][0].get("base64") or result["data"][0].get("url")
            if image_url:
                if image_url.startswith("data:"):
                    # Base64 encoded
                    import base64
                    image_data = base64.b64decode(image_url.split(",")[1])
                else:
                    # URL
                    img_response = requests.get(image_url, timeout=30)
                    image_data = img_response.content

                with open(output_path, "wb") as f:
                    f.write(image_data)
                print(f"Logo saved to: {output_path}")
                return str(output_path)

        # Need to poll for result
        task_id = result.get("task_id")
        if task_id:
            print(f"Task created: {task_id}. Polling for completion...")

            for _ in range(30):  # Max 60 seconds
                time.sleep(2)
                status_response = requests.get(
                    f"https://api.freepik.com/v1/ai/mystic/{task_id}",
                    headers={"x-freepik-api-key": api_key},
                    timeout=30
                )
                status = status_response.json()

                if status.get("status") == "COMPLETED":
                    image_url = status.get("generated", [None])[0]
                    if image_url:
                        img_response = requests.get(image_url, timeout=30)
                        with open(output_path, "wb") as f:
                            f.write(img_response.content)
                        print(f"Logo saved to: {output_path}")
                        return str(output_path)
                elif status.get("status") == "FAILED":
                    print(f"Logo generation failed: {status}")
                    return None

                print(f"  Status: {status.get('status', 'unknown')}")

    except Exception as e:
        print(f"Logo generation error: {e}")

    return None


# =============================================================================
# CONTENT - Expanded from reimagined-brand-strategy.md
# =============================================================================

EXPANDED_CONTENT = {
    'vision': {
        'headline': "Vision Statement",
        'intro': """We envision a future where AI empowers every educator and student to reach their full potential, transforming education through innovation, accessibility, and lifelong learning.""",
        'expanded': """ReimagineED is more than a brand—it's a movement. We position Black and Latino educators at the forefront of the AI revolution in education, not as followers, but as pioneering leaders reshaping the future of learning.

Our vision challenges the status quo. We reject the notion that AI in education is something that happens to communities of color. Instead, we envision a future where these communities lead the transformation, bringing unique perspectives, cultural wisdom, and innovative thinking to the forefront of educational technology.

As The Disruptor, we push boundaries of what's possible with technology in classrooms. We make waves by centering marginalized voices in the AI conversation. We refuse safe, comfortable narratives about educational technology.""",
        'pillars': [
            ("Disruption as Responsibility", "We disrupt because the status quo isn't serving our students. Transformation isn't optional—it's our duty."),
            ("Technology as Liberation", "AI and tech should remove barriers, not create them. We center equity in every innovation discussion."),
            ("Educators as Innovators", "Black and Latino educators aren't passive recipients of tech—we're the architects of educational futures."),
        ]
    },

    'mission': {
        'headline': "Mission Statement",
        'intro': """Our mission is to empower educators and educational leaders with AI-driven solutions that enhance teaching effectiveness, streamline administrative processes, and create personalized learning experiences for all students.""",
        'pillars': [
            ("Educate", "Provide comprehensive AI training and professional development that honors cultural context and practical classroom realities."),
            ("Employ", "Create career pathways and employment opportunities in AI-enhanced education, ensuring our communities benefit economically from the AI revolution."),
            ("Empower", "Equip educators with tools and knowledge to transform their practice while maintaining their authentic voice and cultural identity."),
        ],
        'additional': [
            ("Community as Catalyst", "Collective wisdom and shared experience drive our approach. We rise together, learning from each other's successes and challenges."),
            ("Excellence Without Exception", "We hold ourselves to the highest standards because our communities deserve nothing less. Mediocrity is not an option."),
        ]
    },

    'thought_leadership': {
        'headline': "Thought Leadership",
        'intro': """ReimagineED leads the national conversation on AI in education through podcasts, webinars, speaking engagements, and strategic content. We don't just participate in the discourse—we shape it.""",
        'podcast': {
            'format': "45-60 minute episodes (edited to 35-45 for pacing)",
            'structure': "News roundup → Deep dive → Guest expert → Actionable takeaways",
            'production': "High-quality audio with dynamic editing—not static conversation",
            'visual': "YouTube video versions with tech-forward motion graphics"
        },
        'content_pillars': [
            ("AI Myth Busting", "Debunking common misconceptions about AI in education"),
            ("Educator Spotlights", "Showcasing innovators in action across K-12 and higher ed"),
            ("Tech Tutorials", "Practical AI tools for classrooms that teachers can use tomorrow"),
            ("Industry Analysis", "What's happening in ed-tech and why it matters for our communities"),
            ("Community Conversations", "Amplifying educator voices and lived experiences"),
        ],
        'platforms': {
            'LinkedIn': "Thought leadership, professional discourse, superintendent-level engagement",
            'Instagram': "Behind-the-scenes, community building, visual quotes",
            'Twitter/X': "Real-time commentary, news curation, debate",
            'TikTok': "Educational tech tips, myth-busting, accessibility",
        }
    },

    'voice_tone': {
        'headline': "Voice & Tone",
        'intro': """The ReimagineED voice is unmistakable. We speak with authority earned through expertise and lived experience. Our tone adapts to context while maintaining our core identity.""",
        'characteristics': [
            ("Bold & Unapologetic", "We don't hedge or soften our message. When we have something to say, we say it directly."),
            ("Intellectually Rigorous", "Grounded in research, not hype. We back our claims with evidence and expertise."),
            ("Culturally Authentic", "Speaks from and to Black and Latino educator experience. We don't code-switch our identity."),
            ("Future-Focused", "Always looking ahead, never backwards. Yesterday's solutions don't solve tomorrow's challenges."),
            ("Action-Oriented", "Disruption requires movement, not just talk. Every piece of content drives toward action."),
        ],
        'do': [
            "Challenge conventional wisdom",
            "Provoke thought and inspire action",
            "Question assumptions about AI and education",
            "Center community voice in every discussion",
        ],
        'dont': [
            "Patronize or oversimplify complex issues",
            "Follow trends blindly",
            "Play it safe with lukewarm takes",
            "Dilute our message for comfort",
        ],
        'example': '"We\'re not asking permission to lead the AI revolution in education—we\'re already here."'
    },

    'brand_applications': {
        'headline': "Brand Applications",
        'intro': """Consistent brand application builds recognition and trust. These guidelines ensure ReimagineED maintains its distinctive identity across all touchpoints.""",
        'applications': [
            ("Podcast Episode Artwork", "Navy background, gold accent typography, consistent episode number placement"),
            ("Social Media Templates", "Platform-optimized sizes, brand colors, clear hierarchy"),
            ("Presentation Decks", "Clean layouts, generous whitespace, hero imagery"),
            ("Email Newsletters", "Mobile-first design, scannable format, clear CTAs"),
            ("Website Elements", "Responsive design, accessibility-first, fast loading"),
            ("Event Materials", "Print-ready formats, QR codes for digital connection"),
        ]
    },

    'target_audience': {
        'primary': {
            'title': "Educational Change Agents",
            'demographics': "Black and Latino teachers, principals, superintendents, instructional coaches",
            'age': "28-55",
            'traits': [
                "Tech-savvy, forward-thinking",
                "Frustrated with status quo",
                "Want to lead change, not follow trends",
                "Seek community with others navigating AI integration",
            ],
            'values': "Excellence, equity, innovation, cultural authenticity",
            'pain_points': "Excluded from mainstream AI ed-tech conversations, lack of culturally relevant resources",
            'aspirations': "Lead districts/schools through AI transformation, become thought leaders",
        },
        'secondary': [
            "College administrators and workforce development professionals",
            "Policymakers and education technology vendors",
            "Corporate partners seeking authentic educational partnerships",
        ]
    }
}


# =============================================================================
# WORD DOCUMENT GENERATOR
# =============================================================================

class ReimagineEDBrandGuide:
    """Generates the ReimagineED Brand Style Guide Word document."""

    def __init__(self):
        self.doc = Document()
        self.extracted_images = {}
        self.logo_path = None
        self._setup_page_layout()
        self._setup_styles()

    def _setup_page_layout(self):
        """Configure page margins and size."""
        for section in self.doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _setup_styles(self):
        """Create custom paragraph styles."""
        styles = self.doc.styles

        # Cover Title
        try:
            cover_title = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            cover_title = styles['CoverTitle']
        cover_title.font.name = 'Arial'
        cover_title.font.size = Pt(48)
        cover_title.font.bold = True
        cover_title.font.color.rgb = COLORS['navy']
        cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.space_after = Pt(6)

        # Section Header
        try:
            section_header = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            section_header = styles['SectionHeader']
        section_header.font.name = 'Arial'
        section_header.font.size = Pt(28)
        section_header.font.bold = True
        section_header.font.color.rgb = COLORS['navy']
        section_header.paragraph_format.space_before = Pt(0)
        section_header.paragraph_format.space_after = Pt(18)

        # Subheader
        try:
            subheader = styles.add_style('Subheader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            subheader = styles['Subheader']
        subheader.font.name = 'Arial'
        subheader.font.size = Pt(16)
        subheader.font.bold = True
        subheader.font.color.rgb = COLORS['electric_blue']
        subheader.paragraph_format.space_before = Pt(18)
        subheader.paragraph_format.space_after = Pt(8)

        # Body Text
        try:
            body = styles.add_style('BrandBody', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body = styles['BrandBody']
        body.font.name = 'Arial'
        body.font.size = Pt(11)
        body.font.color.rgb = COLORS['charcoal']
        body.paragraph_format.space_after = Pt(10)
        body.paragraph_format.line_spacing = 1.4

        # Pillar Title
        try:
            pillar = styles.add_style('PillarTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            pillar = styles['PillarTitle']
        pillar.font.name = 'Arial'
        pillar.font.size = Pt(12)
        pillar.font.bold = True
        pillar.font.color.rgb = COLORS['navy']
        pillar.paragraph_format.space_before = Pt(12)
        pillar.paragraph_format.space_after = Pt(4)

    def _add_page_header(self, text: str):
        """Add section header."""
        para = self.doc.add_paragraph(text, style='SectionHeader')
        return para

    def _add_subheader(self, text: str):
        """Add subheader."""
        para = self.doc.add_paragraph(text, style='Subheader')
        return para

    def _add_body(self, text: str):
        """Add body text."""
        para = self.doc.add_paragraph(text, style='BrandBody')
        return para

    def _add_pillar(self, title: str, description: str):
        """Add a pillar with title and description."""
        title_para = self.doc.add_paragraph(title, style='PillarTitle')
        desc_para = self.doc.add_paragraph(description, style='BrandBody')
        return title_para, desc_para

    def _add_body_with_label(self, label: str, value: str):
        """Add body text with bold label (no markdown)."""
        para = self.doc.add_paragraph(style='BrandBody')
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(value)
        return para

    def _add_image(self, image_path: str, width: float = 6.0, caption: str = None):
        """Add image with optional caption."""
        if image_path and os.path.exists(image_path):
            try:
                self.doc.add_picture(image_path, width=Inches(width))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if caption:
                    cap_para = self.doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap_para.add_run(caption)
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = COLORS['charcoal']
                return True
            except Exception as e:
                print(f"Error adding image {image_path}: {e}")
        return False

    # =========================================================================
    # PAGE GENERATORS
    # =========================================================================

    def create_cover_page(self):
        """Page 1: Cover page with branding."""
        # Add spacing at top
        for _ in range(3):
            self.doc.add_paragraph()

        # Brand name
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "REIMAGINE" in navy
        run1 = title.add_run("REIMAGINE")
        run1.font.name = 'Arial'
        run1.font.size = Pt(52)
        run1.font.bold = True
        run1.font.color.rgb = COLORS['navy']

        # "ED" in gold
        run2 = title.add_run("ED")
        run2.font.name = 'Arial'
        run2.font.size = Pt(52)
        run2.font.bold = True
        run2.font.color.rgb = COLORS['gold']

        # Tagline
        tagline = self.doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tagline.add_run("The Disruptor in AI Education")
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.color.rgb = COLORS['electric_blue']

        self.doc.add_paragraph()

        # Add logo image if available
        if self.logo_path and os.path.exists(self.logo_path):
            self._add_image(self.logo_path, width=5.0)
            self.doc.add_paragraph()

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Brand Style Guide")
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLORS['navy']

        # Add spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Footer info
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("The Right Path Educational Consulting Inc.")
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS['charcoal']

        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run("Version 2.0 | 2025")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['charcoal']

        self.doc.add_page_break()

    def create_toc_page(self):
        """Page 2: Table of Contents."""
        self._add_page_header("Table of Contents")

        toc_items = [
            ("1. Vision Statement", "3"),
            ("2. Mission Statement", "4"),
            ("3. Brand Evolution", "5"),
            ("4. Thought Leadership", "6"),
            ("5. Logo System", "7"),
            ("6. Color Palette", "8"),
            ("7. Typography System", "9"),
            ("8. Imagery Guidelines", "10"),
            ("9. Voice & Tone", "11"),
            ("10. Visual Language", "12"),
            ("11. Brand Applications", "13"),
            ("12. Do's & Don'ts", "14"),
            ("13. Social Media Guidelines", "15"),
        ]

        # Create styled table for TOC
        table = self.doc.add_table(rows=len(toc_items), cols=2)
        table.autofit = False

        for i, (title, page) in enumerate(toc_items):
            row = table.rows[i]

            # Title cell
            title_cell = row.cells[0]
            title_cell.width = Inches(5.0)
            p = title_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(title)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = COLORS['navy']

            # Page number cell
            page_cell = row.cells[1]
            page_cell.width = Inches(1.5)
            p = page_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(f"Page {page}")
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS['charcoal']

        self.doc.add_page_break()

    def create_vision_page(self):
        """Page 3: Vision Statement with expanded content."""
        content = EXPANDED_CONTENT['vision']

        self._add_page_header(content['headline'])

        # Hero image if available
        if 2 in self.extracted_images and self.extracted_images[2]:
            self._add_image(self.extracted_images[2][0], width=6.0)
            self.doc.add_paragraph()

        # Intro
        self._add_body(content['intro'])

        # Expanded vision
        self._add_subheader("Our Vision")
        for paragraph in content['expanded'].split('\n\n'):
            if paragraph.strip():
                self._add_body(paragraph.strip())

        # Brand Pillars preview
        self._add_subheader("Core Principles")
        for title, desc in content['pillars']:
            self._add_pillar(title, desc)

        self.doc.add_page_break()

    def create_mission_page(self):
        """Page 4: Mission Statement with Three Pillars."""
        content = EXPANDED_CONTENT['mission']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Three Pillars")
        for title, desc in content['pillars']:
            self._add_pillar(title, desc)

        self._add_subheader("Supporting Values")
        for title, desc in content['additional']:
            self._add_pillar(title, desc)

        self.doc.add_page_break()

    def create_brand_evolution_page(self):
        """Page 5: Brand Evolution."""
        self._add_page_header("Brand Evolution")

        # Hero image if available
        if 4 in self.extracted_images and self.extracted_images[4]:
            self._add_image(self.extracted_images[4][0], width=6.0)
            self.doc.add_paragraph()

        self._add_body("""Our brand represents the evolution of education from traditional methods to AI-enhanced learning. We bridge the past, present, and future of education, honoring proven pedagogy while embracing transformative technology.""")

        self._add_subheader("The Disruptor Archetype")
        self._add_body("""ReimagineED embodies The Disruptor brand archetype. We challenge conventions in how AI is discussed in education. We push boundaries of what's possible with technology in classrooms. We make waves by centering marginalized voices in the AI conversation.""")

        self._add_subheader("Our Journey")

        timeline_items = [
            ("Past", "Traditional educational methods served their time but left gaps in equity and accessibility."),
            ("Present", "AI technology emerges as a transformative force, but without intentional leadership, it risks replicating existing inequities."),
            ("Future", "ReimagineED leads a movement where Black and Latino educators architect the future of AI-enhanced learning."),
        ]

        for phase, desc in timeline_items:
            self._add_pillar(phase, desc)

        self.doc.add_page_break()

    def create_thought_leadership_page(self):
        """Page 6: Thought Leadership (major expansion)."""
        content = EXPANDED_CONTENT['thought_leadership']

        self._add_page_header(content['headline'])

        # Hero image if available
        if 6 in self.extracted_images and self.extracted_images[6]:
            self._add_image(self.extracted_images[6][0], width=5.5)
            self.doc.add_paragraph()

        self._add_body(content['intro'])

        # Podcast format
        self._add_subheader("Podcast Format")
        podcast = content['podcast']
        self._add_body_with_label("Format", podcast['format'])
        self._add_body_with_label("Structure", podcast['structure'])
        self._add_body_with_label("Production", podcast['production'])
        self._add_body_with_label("Visual", podcast['visual'])

        # Content pillars
        self._add_subheader("Content Pillars")
        for title, desc in content['content_pillars']:
            self._add_pillar(title, desc)

        # Platform strategy
        self._add_subheader("Platform Strategy")
        for platform, strategy in content['platforms'].items():
            self._add_pillar(platform, strategy)

        self.doc.add_page_break()

    def create_logo_page(self):
        """Page 7: Logo System."""
        self._add_page_header("Logo System")

        self._add_subheader("Primary Logo")
        self._add_body("""The ReimagineED logo combines modern typography with tech-forward sensibility. The wordmark uses bold geometric letterforms that convey innovation and authority.""")

        # Add generated logo if available
        if self.logo_path and os.path.exists(self.logo_path):
            self._add_image(self.logo_path, width=4.5, caption="Primary Logo - Full Color")
            self.doc.add_paragraph()

        self._add_subheader("Logo Construction")
        self._add_body(""""REIMAGINE" appears in Deep Navy Blue (#0B1D3A), representing authority and technological sophistication. "ED" appears in Breakthrough Gold (#FFD33A), symbolizing innovation and excellence.""")

        self._add_subheader("Clear Space")
        self._add_body("""Maintain a minimum clear space around the logo equal to the height of the letter 'E' in the wordmark on all sides. This ensures the logo maintains visual impact and isn't crowded by other elements.""")

        self._add_subheader("Minimum Size")
        self._add_body("""Digital: 120px wide | Print: 1.5 inches wide

Below these sizes, legibility is compromised. For smaller applications, use the abbreviated "RE" mark.""")

        self._add_subheader("Logo Variations")
        variations = [
            ("Full Color", "Primary usage on white or light backgrounds"),
            ("Reversed", "White logo on navy or dark backgrounds"),
            ("Monochrome", "Single color for limited production contexts"),
        ]
        for var, desc in variations:
            self._add_pillar(var, desc)

        self.doc.add_page_break()

    def create_color_palette_page(self):
        """Page 8: Color Palette."""
        self._add_page_header("Brand Color Palette")

        self._add_subheader("Primary Colors")

        # Create color table
        colors_table = [
            ("Deep Navy Blue", "#0B1D3A", "RGB(11, 29, 58)", "Authority, intelligence, technological sophistication"),
            ("Breakthrough Gold", "#FFD33A", "RGB(255, 211, 58)", "Innovation, excellence, breakthrough thinking"),
            ("Electric Blue", "#00D9FF", "RGB(0, 217, 255)", "AI, digital transformation, future-forward"),
        ]

        table = self.doc.add_table(rows=len(colors_table) + 1, cols=4)
        table.style = 'Table Grid'

        # Header row
        headers = ["Color Name", "Hex", "RGB", "Usage"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].add_run(header).bold = True

        # Data rows
        for row_idx, (name, hex_val, rgb, usage) in enumerate(colors_table, 1):
            row = table.rows[row_idx]
            row.cells[0].text = name
            row.cells[1].text = hex_val
            row.cells[2].text = rgb
            row.cells[3].text = usage

        self.doc.add_paragraph()

        self._add_subheader("Secondary Colors")
        secondary_colors = [
            ("Vibrant Purple", "#7B2FFF", "Innovation, creativity, disruption"),
            ("Neon Coral", "#FF6B6B", "Energy, action, bold moves"),
            ("Charcoal", "#333333", "Professional depth, technical precision"),
        ]

        for name, hex_val, usage in secondary_colors:
            self._add_pillar(f"{name} ({hex_val})", usage)

        self._add_subheader("Gradient System")
        self._add_body_with_label("Primary Gradient", "Navy → Electric Blue → Gold (tech hero moments)")
        self._add_body_with_label("Secondary Gradient", "Purple → Coral (innovation emphasis)")
        self._add_body_with_label("Accent Gradient", "Navy → Purple → Electric Blue (AI-inspired backgrounds)")

        self.doc.add_page_break()

    def create_typography_page(self):
        """Page 9: Typography System."""
        self._add_page_header("Typography System")

        self._add_subheader("Primary Typeface")
        self._add_body("""Our typography system balances professionalism with tech-forward energy. We use clean, geometric sans-serif fonts that convey innovation while remaining highly readable.""")

        self._add_pillar("Headlines", "Space Grotesk ExtraBold or Arial Bold, 24-36pt\nBold, geometric letterforms that command attention")

        self._add_pillar("Subheadings", "Space Grotesk Bold or Arial Bold, 18-20pt\nClear hierarchy while maintaining visual flow")

        self._add_pillar("Body Text", "Inter Regular or Arial Regular, 11-12pt\nOptimized for extended reading, 1.4-1.5 line height")

        self._add_subheader("Technical Typography")
        self._add_pillar("Monospace", "Fira Code or Courier New\nUse for technical content, code snippets, and data displays")

        self._add_subheader("Typography Hierarchy")
        self._add_body("Consistent typographic hierarchy ensures content is scannable and accessible:")
        self._add_body_with_label("Level 1 (Page Titles)", "28-36pt, Bold, Navy")
        self._add_body_with_label("Level 2 (Section Headers)", "18-20pt, Bold, Electric Blue")
        self._add_body_with_label("Level 3 (Subsections)", "14-16pt, Bold, Navy")
        self._add_body_with_label("Body Copy", "11-12pt, Regular, Charcoal")
        self._add_body_with_label("Captions/Notes", "9-10pt, Regular or Italic, Charcoal")

        self.doc.add_page_break()

    def create_imagery_page(self):
        """Page 10: Imagery Guidelines."""
        self._add_page_header("Imagery Guidelines")

        # Hero image if available
        if 10 in self.extracted_images and self.extracted_images[10]:
            self._add_image(self.extracted_images[10][0], width=5.5)
            self.doc.add_paragraph()

        self._add_body("""All ReimagineED imagery centers Black and Latino educators as leaders and innovators. Our visual language is tech-forward, dynamic, and authentically representative.""")

        self._add_subheader("Primary Subjects")
        subjects = [
            "Black and Latino educators engaging with AI technology",
            "Futuristic classroom settings with holographic displays",
            "Educators as tech innovators (not just users)",
            "AI-augmented teaching moments",
            "Community collaboration with tech overlay",
        ]
        for subject in subjects:
            para = self.doc.add_paragraph(style='BrandBody')
            para.add_run("• ").bold = True
            para.add_run(subject)

        self._add_subheader("Visual Style")
        style_elements = [
            ("Lighting", "Dramatic, with neon/tech accents (blues, purples, golds)"),
            ("Composition", "Dynamic angles, movement, energy"),
            ("Treatment", "Slight futuristic/sci-fi influence without being fantastical"),
            ("Authenticity", "Real human emotion + tech enhancement"),
        ]
        for element, desc in style_elements:
            self._add_pillar(element, desc)

        self._add_subheader("AI Image Generation Prompt Formula")
        self._add_body(""""Professional photograph of [Black/Latino educator demographic] [action] in [futuristic classroom/tech setting], dramatic lighting with [neon blue/purple/gold] accents, tech-forward aesthetic, cinematic composition, authentic human emotion, 8K quality, photorealistic but slightly futuristic, emphasizing innovation and transformation" """)

        self.doc.add_page_break()

    def create_voice_tone_page(self):
        """Page 11: Voice & Tone (major expansion)."""
        content = EXPANDED_CONTENT['voice_tone']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Voice Characteristics")
        for title, desc in content['characteristics']:
            self._add_pillar(title, desc)

        self._add_subheader("Tone Guidelines")

        # Do's
        do_para = self.doc.add_paragraph()
        do_para.add_run("DO: ").bold = True
        do_para.add_run(", ".join(content['do']))
        do_para.style = 'BrandBody'

        # Don'ts
        dont_para = self.doc.add_paragraph()
        dont_para.add_run("DON'T: ").bold = True
        dont_para.add_run(", ".join(content['dont']))
        dont_para.style = 'BrandBody'

        self._add_subheader("Example Voice")
        example_para = self.doc.add_paragraph()
        example_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = example_para.add_run(content['example'])
        run.font.italic = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLORS['navy']

        self.doc.add_page_break()

    def create_visual_language_page(self):
        """Page 12: Visual Language."""
        self._add_page_header("Visual Language")

        self._add_body("""Our visual language extends beyond photography to include patterns, icons, and graphic elements that reinforce the ReimagineED identity.""")

        self._add_subheader("Graphic Elements")
        elements = [
            ("Circuit Patterns", "Subtle tech-inspired line work suggesting AI/neural networks"),
            ("Gradient Overlays", "Brand color gradients for depth and tech feel"),
            ("Geometric Shapes", "Bold, angular forms representing disruption and innovation"),
            ("Data Visualization", "Clean, modern charts and graphs for research presentation"),
        ]
        for element, desc in elements:
            self._add_pillar(element, desc)

        self._add_subheader("Icon Style")
        self._add_body("""Icons should be:
• Geometric and clean-lined
• 2px stroke weight for consistency
• Navy or Electric Blue on light backgrounds
• White or Gold on dark backgrounds
• Functionally clear before stylistically interesting""")

        self._add_subheader("Layout Principles")
        principles = [
            ("Generous Whitespace", "Let content breathe; avoid cluttered layouts"),
            ("Strong Grid", "8-column grid for digital, 12-column for print"),
            ("Bold Headlines", "Headlines should anchor each page/screen"),
            ("Visual Hierarchy", "Clear distinction between content levels"),
        ]
        for principle, desc in principles:
            self._add_pillar(principle, desc)

        self.doc.add_page_break()

    def create_brand_applications_page(self):
        """Page 13: Brand Applications (expanded)."""
        content = EXPANDED_CONTENT['brand_applications']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Application Guidelines")
        for title, desc in content['applications']:
            self._add_pillar(title, desc)

        self._add_subheader("Digital Specifications")
        specs = [
            ("Website", "Responsive design, accessibility-first, maximum 3-second load time"),
            ("Email", "600px max width, inline CSS, mobile-optimized"),
            ("Social Images", "1200×630px (link preview), 1080×1080px (Instagram square)"),
            ("Video", "16:9 aspect ratio, brand intro/outro templates, captions required"),
        ]
        for spec, desc in specs:
            self._add_pillar(spec, desc)

        self.doc.add_page_break()

    def create_dos_donts_page(self):
        """Page 14: Do's & Don'ts."""
        self._add_page_header("Do's & Don'ts")

        self._add_body("""Maintaining brand integrity requires consistent application. These guidelines protect the ReimagineED identity across all uses.""")

        self._add_subheader("Logo Usage")

        dos = [
            "Use approved logo files from the brand asset library",
            "Maintain minimum clear space requirements",
            "Use appropriate color variations for backgrounds",
            "Scale proportionally (lock aspect ratio)",
        ]

        donts = [
            "Stretch, skew, or distort the logo",
            "Change logo colors outside approved palette",
            "Add effects (shadows, glows, outlines)",
            "Place on busy backgrounds that reduce legibility",
            "Recreate or modify the logo in any way",
        ]

        self._add_pillar("DO", "\n".join([f"• {item}" for item in dos]))
        self._add_pillar("DON'T", "\n".join([f"• {item}" for item in donts]))

        self._add_subheader("Color Usage")
        self._add_pillar("DO", "• Use primary colors for key brand moments\n• Apply gradients for tech-forward contexts\n• Ensure sufficient contrast for accessibility")
        self._add_pillar("DON'T", "• Use colors at reduced opacity except for approved overlays\n• Mix brand colors with non-approved palettes\n• Sacrifice legibility for visual effect")

        self.doc.add_page_break()

    def create_social_media_page(self):
        """Page 15: Social Media Guidelines."""
        self._add_page_header("Social Media Guidelines")

        self._add_body("""ReimagineED maintains an active social presence across platforms. Each channel serves a specific purpose in our content ecosystem.""")

        platforms = [
            ("LinkedIn", "Thought leadership hub", [
                "Long-form articles on AI in education",
                "Research highlights and data insights",
                "Professional discourse and superintendent engagement",
                "Event announcements and recaps",
            ]),
            ("Instagram", "Community & culture", [
                "Behind-the-scenes content",
                "Visual quotes and key takeaways",
                "Educator spotlights (Stories & Reels)",
                "Event coverage and community moments",
            ]),
            ("Twitter/X", "Real-time engagement", [
                "Breaking news and commentary",
                "Thread-based deep dives",
                "Live event coverage",
                "Industry debate and discussion",
            ]),
            ("TikTok", "Accessible education", [
                "Quick AI tips for educators",
                "Myth-busting content",
                "Trending audio with brand spin",
                "Student-facing content",
            ]),
        ]

        for platform, purpose, content_types in platforms:
            self._add_subheader(f"{platform}: {purpose}")
            for content_type in content_types:
                para = self.doc.add_paragraph(style='BrandBody')
                para.add_run("• ").bold = True
                para.add_run(content_type)

        # No page break - this is the last page

    def create_resources_page(self):
        """Page 16: Resources & Contact."""
        self._add_page_header("Resources & Contact")

        self._add_subheader("Brand Asset Library")
        self._add_body("""Access approved logos, templates, and brand materials at:
[Brand Asset Portal URL]

All assets are available in multiple formats (PNG, SVG, PDF) for various applications.""")

        self._add_subheader("Template Library")
        templates = [
            "Presentation deck template (PowerPoint/Google Slides)",
            "Social media templates (Canva/Figma)",
            "Email newsletter template",
            "Podcast episode artwork template",
            "Event flyer template",
        ]
        for template in templates:
            para = self.doc.add_paragraph(style='BrandBody')
            para.add_run("• ").bold = True
            para.add_run(template)

        self._add_subheader("Brand Support")
        self._add_body("""For brand usage questions, approval requests, or asset needs:

**Email:** brand@therightpathed.com
**Website:** www.therightpathed.com

**The Right Path Educational Consulting Inc.**
Transforming education through AI leadership""")

        # Footer
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("© 2025 The Right Path Educational Consulting Inc. All rights reserved.")
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['charcoal']

    # =========================================================================
    # MAIN BUILD METHOD
    # =========================================================================

    def build(self, extracted_images: dict, logo_path: str = None):
        """Build the complete brand guide document."""
        self.extracted_images = extracted_images
        self.logo_path = logo_path

        print("Building document pages...")

        # Page 1: Cover
        print("  Page 1: Cover")
        self.create_cover_page()

        # Page 2: Table of Contents
        print("  Page 2: Table of Contents")
        self.create_toc_page()

        # Page 3: Vision Statement
        print("  Page 3: Vision Statement")
        self.create_vision_page()

        # Page 4: Mission Statement
        print("  Page 4: Mission Statement")
        self.create_mission_page()

        # Page 5: Brand Evolution
        print("  Page 5: Brand Evolution")
        self.create_brand_evolution_page()

        # Page 6: Thought Leadership
        print("  Page 6: Thought Leadership")
        self.create_thought_leadership_page()

        # Page 7: Logo System
        print("  Page 7: Logo System")
        self.create_logo_page()

        # Page 8: Color Palette
        print("  Page 8: Color Palette")
        self.create_color_palette_page()

        # Page 9: Typography
        print("  Page 9: Typography System")
        self.create_typography_page()

        # Page 10: Imagery Guidelines
        print("  Page 10: Imagery Guidelines")
        self.create_imagery_page()

        # Page 11: Voice & Tone
        print("  Page 11: Voice & Tone")
        self.create_voice_tone_page()

        # Page 12: Visual Language
        print("  Page 12: Visual Language")
        self.create_visual_language_page()

        # Page 13: Brand Applications
        print("  Page 13: Brand Applications")
        self.create_brand_applications_page()

        # Page 14: Do's & Don'ts
        print("  Page 14: Do's & Don'ts")
        self.create_dos_donts_page()

        # Page 15: Social Media
        print("  Page 15: Social Media Guidelines")
        self.create_social_media_page()

        # Resources page removed per user request

        print("Document build complete.")

    def save(self, output_path: Path):
        """Save the document."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"Document saved: {output_path}")
        return output_path


# =============================================================================
# PDF EXPORT
# =============================================================================

def export_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Export Word document to PDF."""
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"PDF exported: {pdf_path}")
        return pdf_path
    except ImportError:
        print("docx2pdf not available. Please export PDF manually from Word.")
    except Exception as e:
        print(f"PDF export error: {e}")
        print("Please open the Word document and export to PDF manually.")
    return None


# =============================================================================
# MAIN
# =============================================================================

def main():
    """Main execution flow."""
    print("=" * 60)
    print("ReimagineED Brand Guide Generator")
    print("=" * 60)

    # Ensure directories exist
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    EXTRACTED_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    LOGOS_DIR.mkdir(parents=True, exist_ok=True)

    # Phase 1: Extract images from source PDF
    print("\n[Phase 1] Extracting images from PDF...")
    if PDF_SOURCE.exists():
        extracted_images = extract_images_from_pdf(PDF_SOURCE, EXTRACTED_IMAGES_DIR)
        print(f"Extracted images from {len(extracted_images)} pages")
    else:
        print(f"Warning: Source PDF not found at {PDF_SOURCE}")
        extracted_images = {}

    # Phase 2: Generate logo
    print("\n[Phase 2] Generating logo...")
    logo_path = LOGOS_DIR / "reimagined_wordmark.png"
    generated_logo = generate_logo_freepik(logo_path)

    if not generated_logo:
        # Check for existing logo assets
        existing_logos = list(BASE_DIR.glob("assets/branding-guide/cover*.png"))
        if existing_logos:
            generated_logo = str(existing_logos[0])
            print(f"Using existing asset: {generated_logo}")

    # Phase 3: Build Word document
    print("\n[Phase 3] Building Word document...")
    guide = ReimagineEDBrandGuide()
    guide.build(extracted_images, generated_logo)

    # Save document
    output_docx = OUTPUT_DIR / "ReimagineED_Brand_Guide_v2.docx"
    guide.save(output_docx)

    # Phase 4: Export to PDF
    print("\n[Phase 4] Exporting to PDF...")
    output_pdf = OUTPUT_DIR / "ReimagineED_Brand_Guide_v2.pdf"
    export_to_pdf(output_docx, output_pdf)

    # Summary
    print("\n" + "=" * 60)
    print("GENERATION COMPLETE")
    print("=" * 60)
    print(f"Word Document: {output_docx}")
    print(f"PDF Document:  {output_pdf}")
    print("\nNext steps:")
    print("1. Review the Word document for any formatting adjustments")
    print("2. Add any additional images or graphics as needed")
    print("3. Export final PDF from Word if automatic export failed")

    return output_docx, output_pdf


if __name__ == "__main__":
    main()
