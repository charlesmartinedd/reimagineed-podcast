"""
The Right Path Podcast Brand Guide Generator
Full refresh with white/purple aesthetic, Crimson Pro typography.

Usage:
    python generate_right_path_brand_guide.py
"""

import os
import sys
import time
import requests
from pathlib import Path
from io import BytesIO

# PDF processing
import fitz  # PyMuPDF

# Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from PIL import Image

# =============================================================================
# CONFIGURATION
# =============================================================================

# Paths
BASE_DIR = Path(r"C:\Users\MarieLexisDad\repos\the-right-path-podcast")
PDF_SOURCE = BASE_DIR / "assets" / "branding-guide" / "reimagined" / "ReimagineED_Brand_Guide_v2.pdf"
STRATEGY_MD = BASE_DIR / "docs" / "reimagined-brand-strategy.md"
OUTPUT_DIR = BASE_DIR / "assets" / "branding-guide"
EXTRACTED_IMAGES_DIR = BASE_DIR / "assets" / "branding-guide" / "extracted_images"
LOGOS_DIR = BASE_DIR / "assets"

# Brand Colors - NEW: White/Purple palette
COLORS = {
    'white': RGBColor(255, 255, 255),           # #FFFFFF - Primary background
    'purple': RGBColor(107, 45, 139),           # #6B2D8B - Primary accent
    'light_purple': RGBColor(139, 77, 171),     # #8B4DAB - Secondary accent
    'dark_purple': RGBColor(74, 29, 97),        # #4A1D61 - Dark accent
    'charcoal': RGBColor(44, 44, 44),           # #2C2C2C - Body text
    'gray_600': RGBColor(107, 107, 107),        # #6B6B6B - Secondary text
    'gray_200': RGBColor(229, 229, 229),        # #E5E5E5 - Borders
    'off_white': RGBColor(250, 250, 250),       # #FAFAFA - Subtle background
}

# =============================================================================
# IMAGE EXTRACTION
# =============================================================================

def extract_images_from_pdf(pdf_path: Path, output_dir: Path) -> dict:
    """Extract images from PDF, skipping page 6 (Leadership Team)."""
    output_dir.mkdir(parents=True, exist_ok=True)
    images = {}

    print(f"Extracting images from: {pdf_path}")
    doc = fitz.open(str(pdf_path))

    for page_num in range(len(doc)):
        # Skip Leadership Team page (page 6, index 5)
        if page_num == 5:
            print(f"  Skipping page {page_num + 1} (Leadership Team)")
            continue

        page = doc[page_num]
        image_list = page.get_images()

        if not image_list:
            continue

        images[page_num] = []

        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # Save image
                image_filename = f"page{page_num + 1}_img{img_index}.{image_ext}"
                image_path = output_dir / image_filename

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                images[page_num].append(str(image_path))
                print(f"  Extracted: {image_filename}")
            except Exception as e:
                print(f"  Error extracting image {img_index} from page {page_num + 1}: {e}")

    doc.close()
    return images


# =============================================================================
# CONTENT - Updated for The Right Path Podcast
# =============================================================================

EXPANDED_CONTENT = {
    'vision': {
        'headline': "Vision Statement",
        'intro': """We envision a future where AI empowers every educator and student to reach their full potential, transforming education through innovation, accessibility, and lifelong learning.""",
        'expanded': """The Right Path Podcast is more than a media platform—it's a leadership resource. We guide educators, administrators, and community leaders through the intersection of AI, workforce development, and educational transformation.

Our vision centers humanity and equity in every conversation about technology. We believe AI adoption should be a tool for removing barriers and creating opportunity—not another mechanism that widens gaps.

As The Guide, we illuminate pathways forward. We provide clarity in complexity, practical wisdom for real-world challenges, and a trusted voice for those navigating unprecedented change in education.""",
        'pillars': [
            ("Clarity Through Complexity", "We make AI accessible and understandable. No jargon, no hype—just practical knowledge educators can use."),
            ("Equity as Foundation", "Every innovation discussion centers on access and opportunity. Technology should serve all students."),
            ("Educators as Leaders", "We position educators not as passive recipients of change, but as architects of educational futures."),
        ]
    },

    'mission': {
        'headline': "Mission Statement",
        'intro': """Our mission is to empower educators and educational leaders with AI-driven solutions that enhance teaching effectiveness, streamline administrative processes, and create personalized learning experiences for all students.""",
        'pillars': [
            ("Educate", "Build AI literacy and trust through clarity. Demystify technology for educators, leaders, and communities with practical knowledge they can apply today."),
            ("Employ", "Connect K-12 and higher education to workforce pipelines. Create career pathways that prepare students for the jobs of tomorrow with real-world AI skills."),
            ("Empower", "Amplify voices missing from the AI conversation. Center equity and access to ensure technology serves all students, families, and communities."),
        ],
        'additional': [
            ("Community as Catalyst", "Collective wisdom and shared experience drive our approach. We rise together, learning from each other's successes and challenges."),
            ("Excellence Without Exception", "We hold ourselves to the highest standards because our communities deserve nothing less."),
        ]
    },

    'thought_leadership': {
        'headline': "Thought Leadership",
        'intro': """The Right Path Podcast leads the national conversation on AI in education through podcasts, webinars, speaking engagements, and strategic content. We don't just participate in the discourse—we shape it.""",
        'podcast': {
            'format': "45-minute episodes with three structured segments",
            'structure': "Educate (0-15 min) → Employ (15-30 min) → Empower (30-45 min)",
            'production': "High-quality audio with professional editing and clear pacing",
            'visual': "Clean, editorial aesthetic with purple accent branding"
        },
        'content_pillars': [
            ("AI Literacy", "Building understanding and trust through clear, practical knowledge"),
            ("Workforce Development", "Connecting education to career pathways and real-world opportunities"),
            ("Equity & Access", "Centering voices and perspectives often missing from tech conversations"),
            ("Innovation in Practice", "Showcasing what's actually working in classrooms today"),
            ("Community Voices", "Amplifying educator experiences and lived wisdom"),
        ],
        'platforms': {
            'LinkedIn': "Thought leadership, professional discourse, executive engagement",
            'Instagram': "Behind-the-scenes, community building, visual quotes",
            'Twitter/X': "Real-time commentary, news curation, industry discussion",
            'YouTube': "Full episodes, educational content, extended conversations",
        }
    },

    'voice_tone': {
        'headline': "Voice & Tone",
        'intro': """The Right Path Podcast voice is unmistakable. We speak with authority earned through expertise and lived experience. Our tone adapts to context while maintaining our core identity.""",
        'characteristics': [
            ("Authoritative & Empowering", "We speak with confidence grounded in expertise, always lifting others up."),
            ("Clear & Accessible", "Complex ideas made understandable. We never talk down or oversimplify."),
            ("Equity-Centered", "Every conversation considers who benefits and who might be left behind."),
            ("Future-Focused", "Always looking ahead with practical optimism. Change is opportunity."),
            ("Action-Oriented", "Every piece of content drives toward something educators can do today."),
        ],
        'do': [
            "Illuminate pathways forward",
            "Provide practical, actionable guidance",
            "Center community voice in every discussion",
            "Acknowledge complexity while offering clarity",
        ],
        'dont': [
            "Patronize or oversimplify complex issues",
            "Chase trends without substance",
            "Ignore equity implications",
            "Promise easy solutions to hard problems",
        ],
        'example': '"We illuminate the path forward—helping educators lead the transformation they want to see."'
    },

    'brand_applications': {
        'headline': "Brand Applications",
        'intro': """Consistent brand application builds recognition and trust. These guidelines ensure The Right Path Podcast maintains its distinctive identity across all touchpoints.""",
        'applications': [
            ("Podcast Episode Artwork", "White background, purple accent typography, consistent episode number placement"),
            ("Social Media Templates", "Platform-optimized sizes, white-dominant with purple accents"),
            ("Presentation Decks", "Clean layouts, generous whitespace, editorial photography"),
            ("Email Newsletters", "Mobile-first design, scannable format, clear CTAs"),
            ("Website Elements", "Responsive design, accessibility-first, fast loading"),
            ("Event Materials", "Print-ready formats, QR codes for digital connection"),
        ]
    },

    'target_audience': {
        'primary': {
            'title': "Educational Leaders",
            'demographics': "Principals, superintendents, instructional coaches, district administrators",
            'age': "35-60",
            'traits': [
                "Responsible for leading change in their organizations",
                "Need practical guidance, not theoretical frameworks",
                "Value equity and community impact",
                "Seek trusted voices for AI integration decisions",
            ],
            'values': "Excellence, equity, innovation, community",
            'pain_points': "Overwhelmed by AI hype, lack of practical guidance, equity concerns",
            'aspirations': "Lead successful AI integration while maintaining human-centered values",
        },
        'secondary': [
            "Classroom teachers seeking practical AI tools",
            "Higher education administrators and workforce development professionals",
            "Policymakers and education technology decision-makers",
        ]
    }
}


# =============================================================================
# WORD DOCUMENT GENERATOR
# =============================================================================

class RightPathBrandGuide:
    """Generates The Right Path Podcast Brand Style Guide Word document."""

    def __init__(self):
        self.doc = Document()
        self.extracted_images = {}
        self.logo_path = None
        self._setup_page_layout()
        self._setup_styles()

    def _setup_page_layout(self):
        """Configure page margins and size."""
        for section in self.doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _setup_styles(self):
        """Create custom paragraph styles with new typography."""
        styles = self.doc.styles

        # Cover Title - Crimson Pro style (using Georgia as fallback)
        try:
            cover_title = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            cover_title = styles['CoverTitle']
        cover_title.font.name = 'Georgia'  # Fallback for Crimson Pro
        cover_title.font.size = Pt(48)
        cover_title.font.bold = True
        cover_title.font.color.rgb = COLORS['purple']
        cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.space_after = Pt(6)

        # Section Header - Purple accent
        try:
            section_header = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            section_header = styles['SectionHeader']
        section_header.font.name = 'Georgia'
        section_header.font.size = Pt(28)
        section_header.font.bold = True
        section_header.font.color.rgb = COLORS['purple']
        section_header.paragraph_format.space_before = Pt(0)
        section_header.paragraph_format.space_after = Pt(18)

        # Subheader - Light purple accent
        try:
            subheader = styles.add_style('Subheader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            subheader = styles['Subheader']
        subheader.font.name = 'Georgia'
        subheader.font.size = Pt(16)
        subheader.font.bold = True
        subheader.font.color.rgb = COLORS['purple']
        subheader.paragraph_format.space_before = Pt(18)
        subheader.paragraph_format.space_after = Pt(8)

        # Body Text - DM Sans style (using Calibri as fallback)
        try:
            body = styles.add_style('BrandBody', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body = styles['BrandBody']
        body.font.name = 'Calibri'  # Fallback for DM Sans
        body.font.size = Pt(11)
        body.font.color.rgb = COLORS['charcoal']
        body.paragraph_format.space_after = Pt(10)
        body.paragraph_format.line_spacing = 1.5

        # Pillar Title
        try:
            pillar = styles.add_style('PillarTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            pillar = styles['PillarTitle']
        pillar.font.name = 'Georgia'
        pillar.font.size = Pt(12)
        pillar.font.bold = True
        pillar.font.color.rgb = COLORS['purple']
        pillar.paragraph_format.space_before = Pt(12)
        pillar.paragraph_format.space_after = Pt(4)

    def _add_page_header(self, text: str):
        """Add section header."""
        para = self.doc.add_paragraph(text, style='SectionHeader')
        return para

    def _add_subheader(self, text: str):
        """Add subheader."""
        para = self.doc.add_paragraph(text, style='Subheader')
        return para

    def _add_body(self, text: str):
        """Add body text."""
        para = self.doc.add_paragraph(text, style='BrandBody')
        return para

    def _add_pillar(self, title: str, description: str):
        """Add a pillar with title and description."""
        title_para = self.doc.add_paragraph(title, style='PillarTitle')
        desc_para = self.doc.add_paragraph(description, style='BrandBody')
        return title_para, desc_para

    def _add_body_with_label(self, label: str, value: str):
        """Add body text with bold label."""
        para = self.doc.add_paragraph(style='BrandBody')
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(value)
        return para

    def _add_image(self, image_path: str, width: float = 6.0, caption: str = None):
        """Add image with optional caption."""
        if image_path and os.path.exists(image_path):
            try:
                self.doc.add_picture(image_path, width=Inches(width))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if caption:
                    cap_para = self.doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap_para.add_run(caption)
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = COLORS['gray_600']
                return True
            except Exception as e:
                print(f"Error adding image {image_path}: {e}")
        return False

    # =========================================================================
    # PAGE GENERATORS
    # =========================================================================

    def create_cover_page(self):
        """Page 1: Cover page with new branding."""
        # Add spacing at top
        for _ in range(3):
            self.doc.add_paragraph()

        # Brand name - "THE RIGHT PATH" in purple
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = title.add_run("THE RIGHT PATH")
        run1.font.name = 'Georgia'
        run1.font.size = Pt(48)
        run1.font.bold = True
        run1.font.color.rgb = COLORS['purple']

        # "PODCAST" on second line
        podcast_line = self.doc.add_paragraph()
        podcast_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = podcast_line.add_run("PODCAST")
        run2.font.name = 'Georgia'
        run2.font.size = Pt(36)
        run2.font.bold = True
        run2.font.color.rgb = COLORS['charcoal']

        # Tagline
        tagline = self.doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tagline.add_run("AI in Action. Educate. Employ. Empower.")
        run.font.name = 'Georgia'
        run.font.size = Pt(18)
        run.font.italic = True
        run.font.color.rgb = COLORS['purple']

        self.doc.add_paragraph()

        # Add logo image if available
        if self.logo_path and os.path.exists(self.logo_path):
            self._add_image(self.logo_path, width=2.0)
            self.doc.add_paragraph()

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Brand Style Guide")
        run.font.name = 'Georgia'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLORS['charcoal']

        # Add spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Footer info
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("The Right Path Educational Consulting Inc.")
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS['gray_600']

        version = self.doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run("Version 3.0 | 2025")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['gray_600']

        self.doc.add_page_break()

    def create_toc_page(self):
        """Page 2: Table of Contents."""
        self._add_page_header("Table of Contents")

        toc_items = [
            ("1. Vision Statement", "3"),
            ("2. Mission Statement", "5"),
            ("3. Brand Evolution", "7"),
            ("4. Thought Leadership", "9"),
            ("5. Logo System", "11"),
            ("6. Color Palette", "13"),
            ("7. Typography System", "15"),
            ("8. Imagery Guidelines", "17"),
            ("9. Voice & Tone", "19"),
            ("10. Visual Language", "20"),
            ("11. Brand Applications", "21"),
            ("12. Do's & Don'ts", "22"),
        ]

        # Create styled table for TOC
        table = self.doc.add_table(rows=len(toc_items), cols=2)
        table.autofit = False

        for i, (title, page) in enumerate(toc_items):
            row = table.rows[i]

            # Title cell
            title_cell = row.cells[0]
            title_cell.width = Inches(5.0)
            p = title_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(title)
            run.font.name = 'Georgia'
            run.font.size = Pt(12)
            run.font.color.rgb = COLORS['charcoal']

            # Page number cell
            page_cell = row.cells[1]
            page_cell.width = Inches(1.5)
            p = page_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(f"Page {page}")
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS['gray_600']

        self.doc.add_page_break()

    def create_vision_page(self):
        """Page 3-4: Vision Statement with expanded content."""
        content = EXPANDED_CONTENT['vision']

        self._add_page_header(content['headline'])

        # Intro
        self._add_body(content['intro'])

        # Expanded vision
        self._add_subheader("Our Vision")
        for paragraph in content['expanded'].split('\n\n'):
            if paragraph.strip():
                self._add_body(paragraph.strip())

        # Core Principles
        self._add_subheader("Core Principles")
        for title, desc in content['pillars']:
            self._add_pillar(title, desc)

        self.doc.add_page_break()

    def create_mission_page(self):
        """Page 5-6: Mission Statement with Three Pillars."""
        content = EXPANDED_CONTENT['mission']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Three Pillars: Educate. Employ. Empower.")
        for title, desc in content['pillars']:
            self._add_pillar(f"📌 {title}", desc)

        self._add_subheader("Supporting Values")
        for title, desc in content['additional']:
            self._add_pillar(title, desc)

        self.doc.add_page_break()

    def create_brand_evolution_page(self):
        """Page 7-8: Brand Evolution."""
        self._add_page_header("Brand Evolution")

        self._add_body("""Our brand represents the evolution of education from traditional methods to AI-enhanced learning. We bridge the past, present, and future of education, honoring proven pedagogy while embracing transformative technology.""")

        self._add_subheader("The Guide Archetype")
        self._add_body("""The Right Path Podcast embodies The Guide brand archetype. We illuminate pathways forward for educators navigating change. We provide clarity in complexity, practical wisdom for real-world challenges, and a trusted voice for those leading transformation.""")

        self._add_subheader("Our Journey")

        timeline_items = [
            ("Past", "Traditional educational methods served their time but left gaps in equity and accessibility."),
            ("Present", "AI technology emerges as a transformative force. With intentional leadership, it can close gaps rather than widen them."),
            ("Future", "The Right Path Podcast guides educators to architect the future of AI-enhanced learning with humanity at the center."),
        ]

        for phase, desc in timeline_items:
            self._add_pillar(phase, desc)

        self.doc.add_page_break()

    def create_thought_leadership_page(self):
        """Page 9-10: Thought Leadership."""
        content = EXPANDED_CONTENT['thought_leadership']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        # Podcast format
        self._add_subheader("Podcast Format")
        podcast = content['podcast']
        self._add_body_with_label("Format", podcast['format'])
        self._add_body_with_label("Structure", podcast['structure'])
        self._add_body_with_label("Production", podcast['production'])
        self._add_body_with_label("Visual", podcast['visual'])

        # Content pillars
        self._add_subheader("Content Pillars")
        for title, desc in content['content_pillars']:
            self._add_pillar(title, desc)

        # Platform strategy
        self._add_subheader("Platform Strategy")
        for platform, strategy in content['platforms'].items():
            self._add_pillar(platform, strategy)

        self.doc.add_page_break()

    def create_logo_page(self):
        """Page 11-12: Logo System."""
        self._add_page_header("Logo System")

        self._add_subheader("Primary Logo")
        self._add_body("""The Right Path Podcast logo features the purple R-in-circle mark. This clean, modern symbol represents guidance, direction, and the path forward.""")

        # Add logo if available
        if self.logo_path and os.path.exists(self.logo_path):
            self._add_image(self.logo_path, width=2.5, caption="Primary Logo - Purple Mark")
            self.doc.add_paragraph()

        self._add_subheader("Logo Construction")
        self._add_body("""The R mark uses our primary purple (#6B2D8B), representing wisdom, leadership, and innovation. The circular container suggests completeness and the journey of continuous learning.""")

        self._add_subheader("Clear Space")
        self._add_body("""Maintain a minimum clear space around the logo equal to the height of the 'R' on all sides. This ensures the logo maintains visual impact and isn't crowded by other elements.""")

        self._add_subheader("Minimum Size")
        self._add_body("""Digital: 36px wide | Print: 0.5 inches wide

Below these sizes, use the full wordmark instead.""")

        self._add_subheader("Logo Variations")
        variations = [
            ("Full Color", "Purple mark on white or light backgrounds"),
            ("Reversed", "White mark on purple or dark backgrounds"),
            ("Monochrome", "Single color for limited production contexts"),
        ]
        for var, desc in variations:
            self._add_pillar(var, desc)

        self.doc.add_page_break()

    def create_color_palette_page(self):
        """Page 13-14: Color Palette - NEW white/purple palette."""
        self._add_page_header("Brand Color Palette")

        self._add_body("""Our color palette centers on white backgrounds with purple accents—clean, professional, and editorial. This approach creates visual clarity and lets content breathe.""")

        self._add_subheader("Primary Colors")

        # Create color table
        colors_table = [
            ("White", "#FFFFFF", "RGB(255, 255, 255)", "Primary background, clarity, openness"),
            ("Royal Purple", "#6B2D8B", "RGB(107, 45, 139)", "Primary accent, wisdom, leadership"),
            ("Charcoal", "#2C2C2C", "RGB(44, 44, 44)", "Body text, professionalism"),
        ]

        table = self.doc.add_table(rows=len(colors_table) + 1, cols=4)
        table.style = 'Table Grid'

        # Header row
        headers = ["Color Name", "Hex", "RGB", "Usage"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].add_run(header).bold = True

        # Data rows
        for row_idx, (name, hex_val, rgb, usage) in enumerate(colors_table, 1):
            row = table.rows[row_idx]
            row.cells[0].text = name
            row.cells[1].text = hex_val
            row.cells[2].text = rgb
            row.cells[3].text = usage

        self.doc.add_paragraph()

        self._add_subheader("Secondary Colors")
        secondary_colors = [
            ("Light Purple", "#8B4DAB", "Hover states, highlights, secondary accents"),
            ("Dark Purple", "#4A1D61", "Dark mode, emphasis, contrast"),
            ("Gray 600", "#6B6B6B", "Secondary text, captions"),
            ("Gray 200", "#E5E5E5", "Borders, dividers, subtle backgrounds"),
        ]

        for name, hex_val, usage in secondary_colors:
            self._add_pillar(f"{name} ({hex_val})", usage)

        self._add_subheader("Color Ratios")
        self._add_body("""Maintain an 80-90% white background ratio across all materials. Purple should accent, not dominate. This creates the clean, editorial aesthetic that defines our brand.""")

        self.doc.add_page_break()

    def create_typography_page(self):
        """Page 15-16: Typography System - NEW Crimson Pro/DM Sans."""
        self._add_page_header("Typography System")

        self._add_body("""Our typography pairs editorial elegance with modern clarity. Crimson Pro brings gravitas to headlines; DM Sans ensures readability in body text.""")

        self._add_subheader("Primary Typefaces")

        self._add_pillar("Headlines: Crimson Pro", "Elegant serif with editorial character. Use Bold (700) for main headlines, SemiBold (600) for subheadings. Sizes: 28-48pt for headlines, 18-24pt for subheads.")

        self._add_pillar("Body Text: DM Sans", "Clean, highly readable sans-serif. Use Regular (400) for body, Medium (500) for emphasis. Size: 11-14pt with 1.5-1.7 line height.")

        self._add_subheader("Typography Hierarchy")
        self._add_body("Consistent hierarchy ensures content is scannable and accessible:")
        self._add_body_with_label("Level 1 (Page Titles)", "Crimson Pro Bold, 28-36pt, Purple")
        self._add_body_with_label("Level 2 (Section Headers)", "Crimson Pro SemiBold, 18-22pt, Purple")
        self._add_body_with_label("Level 3 (Subsections)", "Crimson Pro SemiBold, 14-16pt, Charcoal")
        self._add_body_with_label("Body Copy", "DM Sans Regular, 11-14pt, Charcoal")
        self._add_body_with_label("Captions", "DM Sans Regular, 9-10pt, Gray 600")

        self._add_subheader("Web Fallbacks")
        self._add_body("""When Crimson Pro is unavailable, use Georgia. When DM Sans is unavailable, use system sans-serif or Calibri.""")

        self.doc.add_page_break()

    def create_imagery_page(self):
        """Page 17-18: Imagery Guidelines."""
        self._add_page_header("Imagery Guidelines")

        self._add_body("""All imagery should feel professional, authentic, and human-centered. We feature real educators in genuine moments of engagement, learning, and leadership.""")

        self._add_subheader("Photography Style")
        style_elements = [
            ("Subjects", "Educators, administrators, students engaged in meaningful work"),
            ("Lighting", "Clean, professional lighting—bright but not harsh"),
            ("Composition", "Thoughtful framing, human connection, genuine emotion"),
            ("Color Treatment", "Natural colors; avoid heavy filters or artificial effects"),
            ("Background", "Clean, uncluttered; school/professional environments"),
        ]
        for element, desc in style_elements:
            self._add_pillar(element, desc)

        self._add_subheader("Image Do's")
        dos = [
            "Feature diverse representation authentically",
            "Show real engagement with technology",
            "Capture genuine expressions and interactions",
            "Use images that feel current, not dated",
        ]
        for item in dos:
            para = self.doc.add_paragraph(style='BrandBody')
            para.add_run("✓ ").bold = True
            para.add_run(item)

        self._add_subheader("Image Don'ts")
        donts = [
            "Generic stock photos with forced diversity",
            "Outdated technology (overhead projectors, old computers)",
            "Overly posed or artificial scenarios",
            "Heavy filters or neon effects",
        ]
        for item in donts:
            para = self.doc.add_paragraph(style='BrandBody')
            para.add_run("✗ ").bold = True
            para.add_run(item)

        self.doc.add_page_break()

    def create_voice_tone_page(self):
        """Page 19: Voice & Tone."""
        content = EXPANDED_CONTENT['voice_tone']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Voice Characteristics")
        for title, desc in content['characteristics']:
            self._add_pillar(title, desc)

        self._add_subheader("Tone Guidelines")

        # Do's
        do_para = self.doc.add_paragraph()
        do_para.add_run("DO: ").bold = True
        do_para.add_run(" • ".join(content['do']))
        do_para.style = 'BrandBody'

        # Don'ts
        dont_para = self.doc.add_paragraph()
        dont_para.add_run("DON'T: ").bold = True
        dont_para.add_run(" • ".join(content['dont']))
        dont_para.style = 'BrandBody'

        self._add_subheader("Example Voice")
        example_para = self.doc.add_paragraph()
        example_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = example_para.add_run(content['example'])
        run.font.italic = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLORS['purple']

        self.doc.add_page_break()

    def create_visual_language_page(self):
        """Page 20: Visual Language."""
        self._add_page_header("Visual Language")

        self._add_body("""Our visual language extends beyond photography to include patterns, icons, and graphic elements that reinforce The Right Path Podcast identity.""")

        self._add_subheader("Graphic Elements")
        elements = [
            ("Clean Lines", "Simple, purposeful line work—no unnecessary decoration"),
            ("Generous Whitespace", "Let content breathe; avoid cluttered layouts"),
            ("Purple Accents", "Strategic use of purple for emphasis and branding"),
            ("Subtle Backgrounds", "Off-white or very light gray for section differentiation"),
        ]
        for element, desc in elements:
            self._add_pillar(element, desc)

        self._add_subheader("Icon Style")
        self._add_body("""Icons should be:
• Clean and simple (2px stroke weight)
• Purple on white backgrounds
• White on purple backgrounds
• Functionally clear before stylistically interesting""")

        self._add_subheader("Layout Principles")
        principles = [
            ("80-90% White", "Backgrounds should be predominantly white"),
            ("Strong Hierarchy", "Clear distinction between content levels"),
            ("Consistent Spacing", "Use 8px grid system for all spacing"),
            ("Left-Aligned Text", "Prefer left alignment for readability"),
        ]
        for principle, desc in principles:
            self._add_pillar(principle, desc)

        self.doc.add_page_break()

    def create_brand_applications_page(self):
        """Page 21: Brand Applications."""
        content = EXPANDED_CONTENT['brand_applications']

        self._add_page_header(content['headline'])

        self._add_body(content['intro'])

        self._add_subheader("Application Guidelines")
        for title, desc in content['applications']:
            self._add_pillar(title, desc)

        self._add_subheader("Digital Specifications")
        specs = [
            ("Website", "White background, purple accents, accessibility-first"),
            ("Social Images", "1200×630px (link preview), 1080×1080px (Instagram)"),
            ("Podcast Artwork", "3000×3000px, white background, purple typography"),
            ("Video", "16:9 ratio, white/purple intro/outro, captions required"),
        ]
        for spec, desc in specs:
            self._add_pillar(spec, desc)

        self.doc.add_page_break()

    def create_dos_donts_page(self):
        """Page 22: Do's & Don'ts."""
        self._add_page_header("Do's & Don'ts")

        self._add_body("""Maintaining brand integrity requires consistent application. These guidelines protect The Right Path Podcast identity across all uses.""")

        self._add_subheader("Logo Usage")

        dos = [
            "Use approved logo files from the brand asset library",
            "Maintain minimum clear space requirements",
            "Use purple on white or white on purple only",
            "Scale proportionally (lock aspect ratio)",
        ]

        donts = [
            "Stretch, skew, or distort the logo",
            "Add other colors to the logo",
            "Add effects (shadows, glows, gradients)",
            "Place on busy backgrounds",
            "Recreate or modify the logo",
        ]

        self._add_pillar("DO", "\n".join([f"✓ {item}" for item in dos]))
        self._add_pillar("DON'T", "\n".join([f"✗ {item}" for item in donts]))

        self._add_subheader("Color Usage")
        self._add_pillar("DO", "✓ Maintain 80-90% white backgrounds\n✓ Use purple for accents and emphasis\n✓ Ensure WCAG AA contrast for all text")
        self._add_pillar("DON'T", "✗ Use purple as a dominant background color\n✗ Mix brand colors with unapproved palettes\n✗ Sacrifice readability for visual effect")

        # Footer
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("© 2025 The Right Path Educational Consulting Inc. All rights reserved.")
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['gray_600']

    # =========================================================================
    # MAIN BUILD METHOD
    # =========================================================================

    def build(self, extracted_images: dict, logo_path: str = None):
        """Build the complete brand guide document."""
        self.extracted_images = extracted_images
        self.logo_path = logo_path

        print("Building document pages...")

        # Page 1: Cover
        print("  Page 1: Cover")
        self.create_cover_page()

        # Page 2: Table of Contents
        print("  Page 2: Table of Contents")
        self.create_toc_page()

        # Pages 3-4: Vision Statement
        print("  Pages 3-4: Vision Statement")
        self.create_vision_page()

        # Pages 5-6: Mission Statement
        print("  Pages 5-6: Mission Statement")
        self.create_mission_page()

        # Pages 7-8: Brand Evolution
        print("  Pages 7-8: Brand Evolution")
        self.create_brand_evolution_page()

        # Pages 9-10: Thought Leadership
        print("  Pages 9-10: Thought Leadership")
        self.create_thought_leadership_page()

        # Pages 11-12: Logo System
        print("  Pages 11-12: Logo System")
        self.create_logo_page()

        # Pages 13-14: Color Palette
        print("  Pages 13-14: Color Palette")
        self.create_color_palette_page()

        # Pages 15-16: Typography
        print("  Pages 15-16: Typography System")
        self.create_typography_page()

        # Pages 17-18: Imagery Guidelines
        print("  Pages 17-18: Imagery Guidelines")
        self.create_imagery_page()

        # Page 19: Voice & Tone
        print("  Page 19: Voice & Tone")
        self.create_voice_tone_page()

        # Page 20: Visual Language
        print("  Page 20: Visual Language")
        self.create_visual_language_page()

        # Page 21: Brand Applications
        print("  Page 21: Brand Applications")
        self.create_brand_applications_page()

        # Page 22: Do's & Don'ts
        print("  Page 22: Do's & Don'ts")
        self.create_dos_donts_page()

        print("Document build complete.")

    def save(self, output_path: Path):
        """Save the document."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"Document saved: {output_path}")
        return output_path


# =============================================================================
# PDF EXPORT
# =============================================================================

def export_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Export Word document to PDF."""
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"PDF exported: {pdf_path}")
        return pdf_path
    except ImportError:
        print("docx2pdf not available. Please export PDF manually from Word.")
    except Exception as e:
        print(f"PDF export error: {e}")
        print("Please open the Word document and export to PDF manually.")
    return None


# =============================================================================
# MAIN
# =============================================================================

def main():
    """Main execution flow."""
    print("=" * 60)
    print("The Right Path Podcast Brand Guide Generator")
    print("=" * 60)

    # Ensure directories exist
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    EXTRACTED_IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    # Phase 1: Check for existing images (skip extraction for now)
    print("\n[Phase 1] Checking for existing assets...")
    extracted_images = {}

    # Phase 2: Get logo path
    print("\n[Phase 2] Locating logo...")
    logo_path = LOGOS_DIR / "the-right-path-logo.png"
    if logo_path.exists():
        print(f"Found logo: {logo_path}")
    else:
        print(f"Logo not found at {logo_path}")
        logo_path = None

    # Phase 3: Build Word document
    print("\n[Phase 3] Building Word document...")
    guide = RightPathBrandGuide()
    guide.build(extracted_images, str(logo_path) if logo_path else None)

    # Save document
    output_docx = OUTPUT_DIR / "The_Right_Path_Podcast_Brand_Guide_v3.docx"
    guide.save(output_docx)

    # Phase 4: Export to PDF
    print("\n[Phase 4] Exporting to PDF...")
    output_pdf = OUTPUT_DIR / "The_Right_Path_Podcast_Brand_Guide_v3.pdf"
    export_to_pdf(output_docx, output_pdf)

    # Summary
    print("\n" + "=" * 60)
    print("GENERATION COMPLETE")
    print("=" * 60)
    print(f"Word Document: {output_docx}")
    print(f"PDF Document:  {output_pdf}")
    print("\nNext steps:")
    print("1. Review the Word document for any formatting adjustments")
    print("2. Install Crimson Pro and DM Sans fonts if not present")
    print("3. Export final PDF from Word if automatic export failed")

    return output_docx, output_pdf


if __name__ == "__main__":
    main()
